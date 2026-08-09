#!/usr/bin/env python3
"""Generate the polished two-page professional resume in DOCX and PDF.

Design contract: compact_reference_guide with a customer-pack-inspired candidate
masthead. Named resume overrides use US Letter, 0.64-inch margins, Arial,
single-column flow, restrained navy/teal accents, real Word bullets, and no
tables or icons. The one-page ATS resume remains a separate deliverable.
"""

from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DATA = json.loads((ROOT / "data/resume_professional.json").read_text())
EVIDENCE = json.loads((ROOT / "data/evidence.json").read_text())

DOCX_OUT = ROOT / "output/resume/Aaryan-Patel-Professional-Resume.docx"
PDF_OUT = ROOT / "output/pdf/Aaryan-Patel-Professional-Resume.pdf"
PUBLIC_DOCX = ROOT / "public/resume/Aaryan-Patel-Professional-Resume.docx"
PUBLIC_PDF = ROOT / "public/resume/Aaryan-Patel-Professional-Resume.pdf"
for output in (DOCX_OUT, PDF_OUT, PUBLIC_DOCX, PUBLIC_PDF):
    output.parent.mkdir(parents=True, exist_ok=True)

NAVY = "15283B"
TEAL = "0B6477"
MUTED = "506171"
LINE = "B8C7D1"


def set_run_font(run, size=9.5, bold=False, italic=False, color=NAVY):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, *, size=8.3, bold=False, color=TEAL):
    rel = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    props.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    props.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    props.append(size_node)
    if bold:
        props.append(OxmlElement("w:b"))
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=LINE, size="7", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def create_bullet_num_id(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    level.append(fmt)
    marker = OxmlElement("w:lvlText")
    marker.set(qn("w:val"), "\u2022")
    level.append(marker)
    align = OxmlElement("w:lvlJc")
    align.set(qn("w:val"), "left")
    level.append(align)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "270")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_node = OxmlElement("w:numId")
    num_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_node)
    p_pr.append(num_pr)


def configure_docx(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.64)
    section.right_margin = Inches(0.64)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.45)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.05

    styles = document.styles
    section_style = styles.add_style("Professional Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = "Arial"
    section_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    section_style.font.size = Pt(11.1)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor.from_string(TEAL)
    section_style.paragraph_format.space_before = Pt(9)
    section_style.paragraph_format.space_after = Pt(4)
    section_style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AARYAN PATEL  |  PROFESSIONAL RESUME")
    set_run_font(run, 7.2, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Portfolio evidence and reproduction commands: asp53826.github.io")
    set_run_font(run, 7.2, color=MUTED)


def section_heading(document, label):
    paragraph = document.add_paragraph(style="Professional Resume Section")
    run = paragraph.add_run(label.upper())
    set_run_font(run, 11.1, bold=True, color=TEAL)
    add_bottom_border(paragraph)


def role_header(document, left, right, *, size=10.0):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0.4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(left)
    set_run_font(run, size, bold=True)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph.add_run("\t")
    run = paragraph.add_run(right)
    set_run_font(run, size - 0.7, bold=True, color=MUTED)


def add_docx_bullet(document, num_id, text, *, size=9.2, after=2.3):
    paragraph = document.add_paragraph()
    apply_bullet(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.06
    run = paragraph.add_run(text)
    set_run_font(run, size)


def add_contact_line(document, items, *, after=2.5):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    for index, (label, url) in enumerate(items):
        if index:
            run = paragraph.add_run("  |  ")
            set_run_font(run, 8.2, color=MUTED)
        if url:
            add_hyperlink(paragraph, label, url, size=8.2)
        else:
            run = paragraph.add_run(label)
            set_run_font(run, 8.2, color=MUTED)
    return paragraph


def build_docx():
    document = Document()
    configure_docx(document)
    bullet_id = create_bullet_num_id(document)
    contact = DATA["contact"]

    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    run = name.add_run(contact["name"])
    set_run_font(run, 26, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3.5)
    run = title.add_run(contact["title"].upper())
    set_run_font(run, 10.6, bold=True, color=TEAL)

    add_contact_line(document, [
        (contact["location"], None),
        (contact["email"], f'mailto:{contact["email"]}'),
        (contact["linkedinLabel"], contact["linkedinUrl"]),
    ], after=1.2)
    links = add_contact_line(document, [
        ("Portfolio + full project evidence", contact["portfolioUrl"]),
        (contact["githubLabel"], contact["githubUrl"]),
    ], after=4.2)
    add_bottom_border(links, color=TEAL, size="12", space="4")

    section_heading(document, "Professional Profile")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(DATA["profile"])
    set_run_font(run, 9.55)

    section_heading(document, "Core Technical Skills")
    for row in DATA["skills"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.5)
        run = paragraph.add_run(f'{row["label"]}: ')
        set_run_font(run, 9.0, bold=True)
        run = paragraph.add_run(row["value"])
        set_run_font(run, 9.0)

    section_heading(document, "Professional Experience")
    for item in DATA["experience"]:
        role_header(document, item["organization"], item["location"])
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.3)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(item["role"])
        set_run_font(run, 9.2, italic=True, color=MUTED)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)
        paragraph.add_run("\t")
        run = paragraph.add_run(item["dates"])
        set_run_font(run, 8.8, italic=True, color=MUTED)
        if item.get("progression"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1.6)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(item["progression"])
            set_run_font(run, 8.45, italic=True, color=MUTED)
        for bullet in item["bullets"]:
            add_docx_bullet(document, bullet_id, bullet)

    section_heading(document, "Education")
    for item in DATA["education"]:
        role_header(document, item["school"], item["dates"], size=9.8)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.6)
        run = paragraph.add_run(item["degree"])
        set_run_font(run, 9.1, italic=True, color=MUTED)
        if item.get("detail"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(item["detail"])
            set_run_font(run, 8.8)

    document.add_page_break()
    section_heading(document, "Selected Engineering Projects")
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(5)
    run = intro.add_run("Public repositories include source, tests, benchmark commands, and documented failure boundaries.")
    set_run_font(run, 9.2, italic=True, color=MUTED)

    for project in DATA["projects"]:
        header = document.add_paragraph()
        header.paragraph_format.space_before = Pt(3)
        header.paragraph_format.space_after = Pt(1)
        header.paragraph_format.keep_with_next = True
        add_hyperlink(header, project["name"], project["url"], size=10.0, bold=True, color=NAVY)
        run = header.add_run(f'  |  {project["stack"]}')
        set_run_font(run, 8.8, bold=True, color=TEAL)
        for bullet in project["bullets"]:
            add_docx_bullet(document, bullet_id, bullet, size=9.15, after=2.0)

    section_heading(document, "Engineering Practice")
    for text in [
        "Verification: fault injection, differential testing, linearizability, exact-recall baselines, paired seeds, and controlled performance comparisons.",
        "Delivery: Linux development, Git-based workflows, automated tests and GitHub Actions, Docker packaging, documentation, and reproducible commands.",
        "Target roles: systems engineering, software engineering, ML infrastructure, data engineering, database internals, signal processing, and reliability-focused development."
    ]:
        add_docx_bullet(document, bullet_id, text, size=9.15, after=2.0)

    properties = document.core_properties
    properties.title = "Aaryan Patel - Professional Systems and Software Resume"
    properties.subject = "Systems, software, ML infrastructure, and industrial data engineering resume"
    properties.author = "Aaryan Patel"
    properties.keywords = "systems engineer, software engineer, C++, Python, distributed systems, ML infrastructure, data engineering"
    document.save(DOCX_OUT)
    shutil.copy2(DOCX_OUT, PUBLIC_DOCX)


PDF_W, PDF_H = LETTER
PDF_NAVY = HexColor("#15283B")
PDF_TEAL = HexColor("#0B6477")
PDF_MUTED = HexColor("#506171")
PDF_LINE = HexColor("#B8C7D1")
LEFT = 46
RIGHT = PDF_W - 46
WIDTH = RIGHT - LEFT


def clean_text(text):
    return str(text).replace("\u2013", "-").replace("\u2014", "-").replace("\u2011", "-").replace("\u00d7", "x")


def pdf_text(c, text, x, y, font="Helvetica", size=9.2, color=PDF_NAVY):
    value = clean_text(text)
    c.setFillColor(color)
    c.setFont(font, size)
    c.drawString(x, y, value)
    return stringWidth(value, font, size)


def pdf_link(c, text, url, x, y, font="Helvetica", size=8.2, color=PDF_TEAL):
    width = pdf_text(c, text, x, y, font, size, color)
    c.linkURL(url, (x, y - 2, x + width, y + size + 2), relative=0)
    return width


def wrap_lines(text, font, size, width):
    words = clean_text(text).split()
    lines, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and stringWidth(candidate, font, size) > width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def pdf_paragraph(c, text, x, y, width, *, font="Helvetica", size=9.2, leading=11.2, color=PDF_NAVY):
    for line in wrap_lines(text, font, size, width):
        pdf_text(c, line, x, y, font, size, color)
        y -= leading
    return y


def pdf_section(c, label, y):
    y -= 4
    pdf_text(c, label.upper(), LEFT, y, "Helvetica-Bold", 10.8, PDF_TEAL)
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.7)
    c.line(LEFT, y - 4, RIGHT, y - 4)
    return y - 17


def pdf_bullet(c, text, y, *, size=9.0, leading=10.7, after=2.2):
    pdf_text(c, "-", LEFT + 2, y, "Helvetica-Bold", size, PDF_TEAL)
    return pdf_paragraph(c, text, LEFT + 14, y, WIDTH - 14, size=size, leading=leading) - after


def pdf_role(c, organization, location, role, dates, y):
    pdf_text(c, organization, LEFT, y, "Helvetica-Bold", 10.3)
    c.setFillColor(PDF_MUTED)
    c.setFont("Helvetica-Bold", 9.2)
    c.drawRightString(RIGHT, y, clean_text(location))
    y -= 13
    pdf_text(c, role, LEFT, y, "Helvetica-Oblique", 9.35, PDF_MUTED)
    c.setFillColor(PDF_MUTED)
    c.setFont("Helvetica-Oblique", 9.0)
    c.drawRightString(RIGHT, y, clean_text(dates))
    return y - 13


def pdf_header(c, first_page=False):
    contact = DATA["contact"]
    if first_page:
        y = PDF_H - 43
        pdf_text(c, contact["name"], LEFT, y, "Helvetica-Bold", 27, PDF_NAVY)
        y -= 19
        pdf_text(c, contact["title"].upper(), LEFT, y, "Helvetica-Bold", 10.6, PDF_TEAL)
        y -= 16
        parts = [
            (contact["location"], None),
            (contact["email"], f'mailto:{contact["email"]}'),
            (contact["linkedinLabel"], contact["linkedinUrl"]),
        ]
        x = LEFT
        for index, (label, url) in enumerate(parts):
            if index:
                x += pdf_text(c, "  |  ", x, y, size=8.4, color=PDF_MUTED)
            x += pdf_link(c, label, url, x, y, size=8.4) if url else pdf_text(c, label, x, y, size=8.4, color=PDF_MUTED)
        y -= 14
        x = LEFT
        x += pdf_link(c, "Portfolio + full project evidence", contact["portfolioUrl"], x, y, size=8.4)
        x += pdf_text(c, "  |  ", x, y, size=8.4, color=PDF_MUTED)
        pdf_link(c, contact["githubLabel"], contact["githubUrl"], x, y, size=8.4)
        c.setStrokeColor(PDF_TEAL)
        c.setLineWidth(1.1)
        c.line(LEFT, y - 7, RIGHT, y - 7)
        return y - 23
    pdf_text(c, "AARYAN PATEL", LEFT, PDF_H - 33, "Helvetica-Bold", 8.0, PDF_MUTED)
    c.setFillColor(PDF_MUTED)
    c.setFont("Helvetica", 8.0)
    c.drawRightString(RIGHT, PDF_H - 33, "SYSTEMS, SOFTWARE & ML INFRASTRUCTURE")
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.6)
    c.line(LEFT, PDF_H - 39, RIGHT, PDF_H - 39)
    return PDF_H - 59


def pdf_footer(c, page_number):
    text = f"Portfolio evidence and reproduction commands: asp53826.github.io  |  Page {page_number} of 2"
    c.setFillColor(PDF_MUTED)
    c.setFont("Helvetica", 7.2)
    c.drawCentredString(PDF_W / 2, 24, text)


def build_pdf():
    c = canvas.Canvas(str(PDF_OUT), pagesize=LETTER, pageCompression=1)
    c.setTitle("Aaryan Patel - Professional Systems and Software Resume")
    c.setAuthor("Aaryan Patel")
    c.setCreator("Aaryan Patel")
    c.setSubject("Systems, software, ML infrastructure, and industrial data engineering resume")
    c.setKeywords("systems engineer, software engineer, C++, Python, distributed systems, ML infrastructure, data engineering")

    y = pdf_header(c, first_page=True)
    y = pdf_section(c, "Professional Profile", y)
    y = pdf_paragraph(c, DATA["profile"], LEFT, y, WIDTH, size=9.9, leading=12.3) - 3

    y = pdf_section(c, "Core Technical Skills", y)
    for row in DATA["skills"]:
        label = f'{row["label"]}: '
        label_width = pdf_text(c, label, LEFT, y, "Helvetica-Bold", 9.2)
        y = pdf_paragraph(c, row["value"], LEFT + label_width, y, WIDTH - label_width, size=9.2, leading=11.2) - 1.4

    y = pdf_section(c, "Professional Experience", y)
    for item in DATA["experience"]:
        y = pdf_role(c, item["organization"], item["location"], item["role"], item["dates"], y)
        if item.get("progression"):
            y = pdf_paragraph(c, item["progression"], LEFT, y, WIDTH, font="Helvetica-Oblique", size=8.8, leading=10.7, color=PDF_MUTED) - 1.5
        for bullet in item["bullets"]:
            y = pdf_bullet(c, bullet, y, size=9.25, leading=11.2, after=2.4)
        y -= 2

    y = pdf_section(c, "Education", y)
    for item in DATA["education"]:
        pdf_text(c, item["school"], LEFT, y, "Helvetica-Bold", 9.9)
        c.setFillColor(PDF_MUTED)
        c.setFont("Helvetica-Bold", 8.9)
        c.drawRightString(RIGHT, y, clean_text(item["dates"]))
        y -= 11
        pdf_text(c, item["degree"], LEFT, y, "Helvetica-Oblique", 9.2, PDF_MUTED)
        y -= 11
        if item.get("detail"):
            y = pdf_paragraph(c, item["detail"], LEFT, y, WIDTH, size=8.9, leading=10.7) - 2
    if y < 40:
        raise RuntimeError(f"professional resume page 1 overflow: {y:.1f}")
    pdf_footer(c, 1)
    c.showPage()

    y = pdf_header(c, first_page=False)
    y = pdf_section(c, "Selected Engineering Projects", y)
    y = pdf_paragraph(c, "Public repositories include source, tests, benchmark commands, and documented failure boundaries.", LEFT, y, WIDTH, font="Helvetica-Oblique", size=9.3, leading=11.2, color=PDF_MUTED) - 3
    for project in DATA["projects"]:
        name_width = pdf_link(c, project["name"], project["url"], LEFT, y, "Helvetica-Bold", 10.1, PDF_NAVY)
        pdf_text(c, f'  |  {project["stack"]}', LEFT + name_width, y, "Helvetica-Bold", 8.7, PDF_TEAL)
        y -= 13
        for bullet in project["bullets"]:
            y = pdf_bullet(c, bullet, y, size=9.2, leading=11.0, after=2.1)
        y -= 3

    y = pdf_section(c, "Engineering Practice", y)
    practices = [
        "Verification: fault injection, differential testing, linearizability, exact-recall baselines, paired seeds, and controlled performance comparisons.",
        "Delivery: Linux development, Git-based workflows, automated tests and GitHub Actions, Docker packaging, documentation, and reproducible commands.",
        "Target roles: systems engineering, software engineering, ML infrastructure, data engineering, database internals, signal processing, and reliability-focused development.",
    ]
    for text in practices:
        y = pdf_bullet(c, text, y, size=9.1, leading=10.9, after=2.1)
    if y < 40:
        raise RuntimeError(f"professional resume page 2 overflow: {y:.1f}")
    pdf_footer(c, 2)
    c.save()
    shutil.copy2(PDF_OUT, PUBLIC_PDF)


def validate_claim_sources():
    evidence = {project["id"]: project for project in EVIDENCE["projects"]}
    for project in DATA["projects"]:
        source = evidence.get(project["id"])
        if source is None:
            raise ValueError(f'professional resume project missing from evidence manifest: {project["id"]}')
        if project["url"] != source["repo"]:
            raise ValueError(f'professional resume URL drift for {project["id"]}')


if __name__ == "__main__":
    validate_claim_sources()
    build_docx()
    build_pdf()
    print(f"wrote {DOCX_OUT}")
    print(f"wrote {PDF_OUT}")
    print(f"wrote {PUBLIC_DOCX}")
    print(f"wrote {PUBLIC_PDF}")
