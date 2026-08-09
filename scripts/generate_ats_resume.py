#!/usr/bin/env python3
"""Generate the ATS-friendly master resume as DOCX and PDF.

Design contract: compact_reference_guide with a restrained, customer-pack-
inspired candidate masthead. Named resume overrides keep the artifact to one
page: 0.58-inch side margins, 0.46-inch top/bottom margins, Arial typography,
single-column flow, no tables, no icons, and explicit real Word bullets.
"""

from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DATA = json.loads((ROOT / "data/resume.json").read_text())
EVIDENCE = json.loads((ROOT / "data/evidence.json").read_text())

DOCX_OUT = ROOT / "output/resume/Aaryan-Patel-Systems-Resume.docx"
PDF_OUT = ROOT / "output/pdf/Aaryan-Patel-Systems-Resume.pdf"
PUBLIC_DOCX = ROOT / "public/resume/Aaryan-Patel-Systems-Resume.docx"
PUBLIC_PDF = ROOT / "public/resume/Aaryan-Patel-Systems-Resume.pdf"
for path in (DOCX_OUT, PDF_OUT, PUBLIC_DOCX, PUBLIC_PDF):
    path.parent.mkdir(parents=True, exist_ok=True)

INK = "172033"
BLUE = "0B5E75"
MUTED = "4E5D6C"
LINE = "B8C5CF"


def set_run_font(run, size=9.2, bold=False, italic=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, *, size=8.4, bold=False, color=BLUE):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    props.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    props.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    props.append(size_node)
    if bold:
        props.append(OxmlElement("w:b"))
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=LINE, size="8", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def create_bullet_num_id(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def configure_docx(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.46)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    styles = document.styles
    if "Resume Section" not in styles:
        section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    else:
        section_style = styles["Resume Section"]
    section_style.font.name = "Arial"
    section_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    section_style.font.size = Pt(10.3)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor.from_string(BLUE)
    section_style.paragraph_format.space_before = Pt(5.5)
    section_style.paragraph_format.space_after = Pt(2.3)
    section_style.paragraph_format.keep_with_next = True


def docx_section(document, label):
    paragraph = document.add_paragraph(style="Resume Section")
    run = paragraph.add_run(label.upper())
    set_run_font(run, 10.3, bold=True, color=BLUE)
    add_bottom_border(paragraph)
    return paragraph


def docx_role_header(document, left, right, *, size=9.6, after=0.4):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    left_run = paragraph.add_run(left)
    set_run_font(left_run, size, bold=True)
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.1), WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph.add_run("\t")
    right_run = paragraph.add_run(right)
    set_run_font(right_run, size - 0.4, bold=True, color=MUTED)
    return paragraph


def build_docx():
    document = Document()
    configure_docx(document)
    bullet_num_id = create_bullet_num_id(document)
    contact = DATA["contact"]

    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(contact["name"])
    set_run_font(run, 21, bold=True, color=INK)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2.5)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(contact["title"].upper())
    set_run_font(run, 9.6, bold=True, color=BLUE)

    links = document.add_paragraph()
    links.paragraph_format.space_after = Pt(4.0)
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [
        (contact["location"], None),
        (contact["email"], f'mailto:{contact["email"]}'),
        (contact["portfolioLabel"], contact["portfolioUrl"]),
        (contact["githubLabel"], contact["githubUrl"]),
        (contact["linkedinLabel"], contact["linkedinUrl"]),
    ]
    for index, (label, url) in enumerate(parts):
        if index:
            separator = links.add_run("  |  ")
            set_run_font(separator, 8.2, color=MUTED)
        if url:
            add_hyperlink(links, label, url, size=8.2)
        else:
            plain = links.add_run(label)
            set_run_font(plain, 8.2, color=MUTED)
    add_bottom_border(links, color=BLUE, size="12", space="4")

    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(3.5)
    summary.paragraph_format.line_spacing = 1.02
    run = summary.add_run(DATA["summary"])
    set_run_font(run, 9.15)

    docx_section(document, "Technical Skills")
    for row in DATA["skills"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.2)
        label_run = paragraph.add_run(f'{row["label"]}: ')
        set_run_font(label_run, 8.7, bold=True)
        value_run = paragraph.add_run(row["value"])
        set_run_font(value_run, 8.7)

    docx_section(document, "Experience")
    for item in DATA["experience"]:
        docx_role_header(document, item["organization"], item["location"])
        role = document.add_paragraph()
        role.paragraph_format.space_after = Pt(1)
        role.paragraph_format.keep_with_next = True
        role_run = role.add_run(item["role"])
        set_run_font(role_run, 8.9, italic=True, color=MUTED)
        role.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_ALIGN_PARAGRAPH.RIGHT)
        role.add_run("\t")
        dates_run = role.add_run(item["dates"])
        set_run_font(dates_run, 8.7, italic=True, color=MUTED)
        if item.get("progression"):
            progression = document.add_paragraph()
            progression.paragraph_format.space_after = Pt(1)
            progression.paragraph_format.keep_with_next = True
            progress_run = progression.add_run(item["progression"])
            set_run_font(progress_run, 7.9, italic=True, color=MUTED)
        for bullet in item["bullets"]:
            paragraph = document.add_paragraph()
            apply_bullet(paragraph, bullet_num_id)
            paragraph.paragraph_format.space_after = Pt(1.4)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(bullet)
            set_run_font(run, 8.55)

    docx_section(document, "Selected Engineering Projects")
    for project in DATA["projects"]:
        header = document.add_paragraph()
        header.paragraph_format.space_after = Pt(0.1)
        header.paragraph_format.keep_with_next = True
        add_hyperlink(header, project["name"], project["url"], size=8.9, bold=True, color=INK)
        stack = header.add_run(f'  |  {project["stack"]}')
        set_run_font(stack, 8.2, bold=True, color=BLUE)
        bullet = document.add_paragraph()
        apply_bullet(bullet, bullet_num_id)
        bullet.paragraph_format.space_after = Pt(1.5)
        bullet.paragraph_format.line_spacing = 1.0
        run = bullet.add_run(project["bullet"])
        set_run_font(run, 8.45)

    docx_section(document, "Education")
    for item in DATA["education"]:
        docx_role_header(document, item["school"], item["dates"], size=9.1, after=0.1)
        degree = document.add_paragraph()
        degree.paragraph_format.space_after = Pt(0.6)
        run = degree.add_run(item["degree"])
        set_run_font(run, 8.55, italic=True, color=MUTED)
        if item.get("detail"):
            detail = document.add_paragraph()
            detail.paragraph_format.space_after = Pt(0)
            run = detail.add_run(item["detail"])
            set_run_font(run, 8.2)

    properties = document.core_properties
    properties.title = "Aaryan Patel - Systems and Software Engineer Resume"
    properties.subject = "ATS-friendly systems, software, ML infrastructure, and data engineering resume"
    properties.author = "Aaryan Patel"
    properties.keywords = "systems engineer, software engineer, C++, Python, distributed systems, ML infrastructure, data engineering"
    document.save(DOCX_OUT)
    shutil.copy2(DOCX_OUT, PUBLIC_DOCX)


PDF_W, PDF_H = LETTER
PDF_INK = HexColor("#172033")
PDF_BLUE = HexColor("#0B5E75")
PDF_MUTED = HexColor("#4E5D6C")
PDF_LINE = HexColor("#B8C5CF")
LEFT = 42
RIGHT = PDF_W - 42
WIDTH = RIGHT - LEFT


def pdf_text(c, text, x, y, font="Helvetica", size=8.5, color=PDF_INK):
    text = str(text).replace("\u2013", "-").replace("\u2014", "-").replace("\u2011", "-")
    c.setFillColor(color)
    c.setFont(font, size)
    c.drawString(x, y, text)
    return stringWidth(text, font, size)


def pdf_link(c, text, url, x, y, font="Helvetica", size=8.2, color=PDF_BLUE):
    width = pdf_text(c, text, x, y, font, size, color)
    c.linkURL(url, (x, y - 2, x + width, y + size + 1), relative=0)
    return width


def wrap_lines(text, font, size, width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and stringWidth(candidate, font, size) > width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def pdf_paragraph(c, text, x, y, width, *, font="Helvetica", size=8.35, leading=10.0, color=PDF_INK):
    for line in wrap_lines(text, font, size, width):
        pdf_text(c, line, x, y, font, size, color)
        y -= leading
    return y


def pdf_section(c, label, y):
    y -= 3
    pdf_text(c, label.upper(), LEFT, y, "Helvetica-Bold", 9.8, PDF_BLUE)
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.65)
    c.line(LEFT, y - 3, RIGHT, y - 3)
    return y - 15


def pdf_role(c, left, right, y, *, size=9.0):
    pdf_text(c, left, LEFT, y, "Helvetica-Bold", size, PDF_INK)
    c.setFont("Helvetica-Bold", size - 0.3)
    c.setFillColor(PDF_MUTED)
    c.drawRightString(RIGHT, y, right)
    return y - 11.5


def pdf_bullet(c, text, y, *, size=8.35, leading=10.0):
    pdf_text(c, "-", LEFT + 2, y, "Helvetica-Bold", size, PDF_BLUE)
    return pdf_paragraph(c, text, LEFT + 13, y, WIDTH - 13, size=size, leading=leading) - 1.2


def build_pdf():
    contact = DATA["contact"]
    c = canvas.Canvas(str(PDF_OUT), pagesize=LETTER, pageCompression=1)
    c.setTitle("Aaryan Patel - Systems and Software Engineer Resume")
    c.setAuthor("Aaryan Patel")
    c.setSubject("ATS-friendly systems, software, ML infrastructure, and data engineering resume")

    y = PDF_H - 34
    name_width = stringWidth(contact["name"], "Helvetica-Bold", 20)
    pdf_text(c, contact["name"], (PDF_W - name_width) / 2, y, "Helvetica-Bold", 20, PDF_INK)
    y -= 15
    title = contact["title"].upper()
    title_width = stringWidth(title, "Helvetica-Bold", 9.2)
    pdf_text(c, title, (PDF_W - title_width) / 2, y, "Helvetica-Bold", 9.2, PDF_BLUE)
    y -= 14

    contact_parts = [
        (contact["location"], None),
        (contact["email"], f'mailto:{contact["email"]}'),
        (contact["portfolioLabel"], contact["portfolioUrl"]),
        (contact["githubLabel"], contact["githubUrl"]),
        (contact["linkedinLabel"], contact["linkedinUrl"]),
    ]
    separator = "  |  "
    total = sum(stringWidth(label, "Helvetica", 7.6) for label, _ in contact_parts)
    total += (len(contact_parts) - 1) * stringWidth(separator, "Helvetica", 7.6)
    x = (PDF_W - total) / 2
    for index, (label, url) in enumerate(contact_parts):
        if index:
            x += pdf_text(c, separator, x, y, "Helvetica", 7.6, PDF_MUTED)
        if url:
            x += pdf_link(c, label, url, x, y, size=7.6)
        else:
            x += pdf_text(c, label, x, y, "Helvetica", 7.6, PDF_MUTED)
    c.setStrokeColor(PDF_BLUE)
    c.setLineWidth(1.0)
    c.line(LEFT, y - 7, RIGHT, y - 7)
    y -= 19

    y = pdf_paragraph(c, DATA["summary"], LEFT, y, WIDTH, size=8.85, leading=10.7) - 2.5

    y = pdf_section(c, "Technical Skills", y)
    for row in DATA["skills"]:
        label = f'{row["label"]}: '
        label_width = pdf_text(c, label, LEFT, y, "Helvetica-Bold", 8.25, PDF_INK)
        y = pdf_paragraph(c, row["value"], LEFT + label_width, y, WIDTH - label_width, size=8.25, leading=9.6) - 0.5

    y = pdf_section(c, "Experience", y)
    for item in DATA["experience"]:
        y = pdf_role(c, item["organization"], item["location"], y)
        pdf_text(c, item["role"], LEFT, y, "Helvetica-Oblique", 8.55, PDF_MUTED)
        c.setFont("Helvetica-Oblique", 8.3)
        c.setFillColor(PDF_MUTED)
        c.drawRightString(RIGHT, y, item["dates"])
        y -= 9.5
        if item.get("progression"):
            pdf_text(c, item["progression"], LEFT, y, "Helvetica-Oblique", 7.65, PDF_MUTED)
            y -= 9.2
        for bullet in item["bullets"]:
            y = pdf_bullet(c, bullet, y)
        y -= 1

    y = pdf_section(c, "Selected Engineering Projects", y)
    for project in DATA["projects"]:
        name_width = pdf_link(c, project["name"], project["url"], LEFT, y, "Helvetica-Bold", 8.7, PDF_INK)
        pdf_text(c, f'  |  {project["stack"]}', LEFT + name_width, y, "Helvetica-Bold", 7.9, PDF_BLUE)
        y -= 9.8
        y = pdf_bullet(c, project["bullet"], y, size=8.05, leading=9.5)

    y = pdf_section(c, "Education", y)
    for item in DATA["education"]:
        y = pdf_role(c, item["school"], item["dates"], y, size=8.8)
        pdf_text(c, item["degree"], LEFT, y, "Helvetica-Oblique", 8.2, PDF_MUTED)
        y -= 9.5
        if item.get("detail"):
            y = pdf_paragraph(c, item["detail"], LEFT, y, WIDTH, size=8.0, leading=9.0) - 0.8

    if y < 28:
        raise RuntimeError(f"resume overflow: final baseline {y:.1f}")
    c.save()
    shutil.copy2(PDF_OUT, PUBLIC_PDF)


def validate_claim_sources():
    evidence_by_id = {project["id"]: project for project in EVIDENCE["projects"]}
    for project in DATA["projects"]:
        if project["id"] not in evidence_by_id:
            raise ValueError(f'resume project missing from evidence manifest: {project["id"]}')
        if project["url"] != evidence_by_id[project["id"]]["repo"]:
            raise ValueError(f'resume URL drift for {project["id"]}')


if __name__ == "__main__":
    validate_claim_sources()
    build_docx()
    build_pdf()
    print(f"wrote {DOCX_OUT}")
    print(f"wrote {PDF_OUT}")
    print(f"wrote {PUBLIC_DOCX}")
    print(f"wrote {PUBLIC_PDF}")
